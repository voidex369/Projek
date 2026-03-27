import os
import uuid
import shutil
from datetime import datetime, timedelta
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import aiofiles

app = FastAPI(title="PDF Tools API", version="1.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_DIR = Path(os.getenv('UPLOAD_DIR', '/app/uploads'))
OUTPUT_DIR = Path(os.getenv('OUTPUT_DIR', '/app/outputs'))
MAX_FILE_SIZE = int(os.getenv('MAX_FILE_SIZE', str(500 * 1024 * 1024)))  # default 50MB
AUTO_DELETE_HOURS = int(os.getenv('AUTO_DELETE_HOURS', '24'))

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Helper: validate file size
async def validate_file_size(file: UploadFile, max_size: int) -> int:
    """Check file size without consuming the stream entirely."""
    content_length = file.headers.get('content-length')
    if content_length:
        size = int(content_length)
        if size > max_size:
            raise HTTPException(status_code=400, detail=f"File too large (max {max_size//(1024*1024)}MB)")
        return size
    # Fallback: read in chunks
    size = 0
    chunk_size = 1024*1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        size += len(chunk)
        if size > max_size:
            await file.seek(0)
            raise HTTPException(status_code=400, detail=f"File too large (max {max_size//(1024*1024)}MB)")
    await file.seek(0)
    return size

# Helper: validate file size
async def validate_file_size(file: UploadFile, max_size: int) -> int:
    """Check file size without consuming the stream entirely."""
    # file.file is SpooledTemporaryFile; we can use seek/tell
    # First, get current position (should be 0)
    await file.seek(0)
    # Try to get size from Content-Length header if available
    content_length = file.headers.get('content-length')
    if content_length:
        size = int(content_length)
        if size > max_size:
            raise HTTPException(status_code=400, detail=f"File too large (max {max_size//(1024*1024)}MB)")
        return size
    # Fallback: read in chunks
    size = 0
    chunk_size = 1024*1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        size += len(chunk)
        if size > max_size:
            await file.seek(0)
            raise HTTPException(status_code=400, detail=f"File too large (max {max_size//(1024*1024)}MB)")
    await file.seek(0)
    return size


class CleanupResponse(BaseModel):
    status: str
    deleted_files: int


@app.get("/health")
async def health_check():
    return {"status": "ok", "timestamp": datetime.utcnow().isoformat()}


@app.post("/api/tools/merge")
async def merge_pdf(
    files: list[UploadFile] = File(...),
    compress: bool = Form(False),
    compression_level: str = Form('medium')
):
    """Merge multiple PDFs into one, optionally compress after merge."""
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="At least 2 PDFs required")

    # Validate all files
    for f in files:
        if f.content_type != 'application/pdf':
            raise HTTPException(status_code=400, detail=f"File {f.filename} is not a PDF")
        await validate_file_size(f, MAX_FILE_SIZE)

    # Generate unique output filename
    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}_merged.pdf"

    try:
        # Placeholder merge logic
        # In production, use PyPDF2 or pdf-lib
        await _merge_pdfs(files, output_path, compress, compression_level)

        # Schedule auto-delete
        _schedule_auto_delete(output_path)

        return FileResponse(
            path=output_path,
            filename="merged.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/tools/split")
async def split_pdf(
    file: UploadFile = File(...),
    strategy: str = Form('all')  # 'all' or 'range'
):
    """Split a PDF into multiple pages."""
    if file.content_type != 'application/pdf':
        raise HTTPException(status_code=400, detail="File must be a PDF")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())

    try:
        # Save uploaded file
        input_path = UPLOAD_DIR / f"{task_id}_input.pdf"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # Placeholder split logic
        # In production, extract pages and create individual PDFs
        pages = await _split_pdf(input_path, strategy)

        # Package as ZIP
        zip_path = OUTPUT_DIR / f"{task_id}_split.zip"
        _create_zip(pages, zip_path)

        # Schedule auto-delete
        _schedule_auto_delete(zip_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=zip_path,
            filename="split_pages.zip",
            media_type="application/zip"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/tools/compress")
async def compress_pdf(
    file: UploadFile = File(...),
    level: str = Form('medium')  # 'low', 'medium', 'high'
):
    """Compress a PDF file."""
    if file.content_type != 'application/pdf':
        raise HTTPException(status_code=400, detail="File must be a PDF")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}_compressed.pdf"

    try:
        # Save uploaded file
        input_path = UPLOAD_DIR / f"{task_id}_input.pdf"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # Placeholder compress logic
        await _compress_pdf(input_path, output_path, level)

        # Schedule auto-delete
        _schedule_auto_delete(output_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_path,
            filename="compressed.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/tools/convert/pdf-to-word")
async def convert_pdf_to_word(
    file: UploadFile = File(...),
    preserve_formatting: bool = Form(True),
    include_images: bool = Form(True)
):
    """Convert PDF to Word document."""
    if file.content_type != 'application/pdf':
        raise HTTPException(status_code=400, detail="File must be a PDF")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}.docx"

    try:
        # Save uploaded file
        input_path = UPLOAD_DIR / f"{task_id}_input.pdf"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # Placeholder conversion logic
        await _convert_pdf_to_word(input_path, output_path, preserve_formatting, include_images)

        # Schedule auto-delete
        _schedule_auto_delete(output_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_path,
            filename="converted.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/status")
async def get_status():
    """Get server status."""
    upload_files = list(UPLOAD_DIR.glob('*'))
    output_files = list(OUTPUT_DIR.glob('*'))
    return {
        "status": "running",
        "uploads": len(upload_files),
        "outputs": len(output_files),
        "uptime": datetime.utcnow().isoformat()
    }


@app.post("/api/cleanup")
async def cleanup_old_files(hours: int = AUTO_DELETE_HOURS):
    """Manually trigger cleanup of old files."""
    cutoff = datetime.utcnow() - timedelta(hours=hours)
    deleted = 0

    for directory in [UPLOAD_DIR, OUTPUT_DIR]:
        for file_path in directory.glob('*'):
            if file_path.is_file():
                mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                if mtime < cutoff:
                    file_path.unlink(missing_ok=True)
                    deleted += 1

    return CleanupResponse(status="cleaned", deleted_files=deleted)


# Background cleanup task
def _schedule_auto_delete(filepath: Path):
    """Mark file for auto-deletion after 24 hours."""
    # In production, use Celery or background task
    # For now, just keep track in a simple way
    pass


# Placeholder implementations (to be replaced with actual processing)
async def _merge_pdfs(files: list[UploadFile], output_path: Path, compress: bool, level: str):
    """Merge PDFs using PyPDF2."""
    from PyPDF2 import PdfWriter, PdfReader
    import io

    writer = PdfWriter()

    for file in files:
        content = await file.read()
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            writer.add_page(page)

    # Write output
    with open(output_path, 'wb') as f:
        writer.write(f)

    if compress:
        # Compress the merged PDF (overwrite)
        await _compress_pdf(output_path, output_path, level)


async def _split_pdf(input_path: Path, strategy: str) -> list[Path]:
    """Split PDF into individual pages."""
    from PyPDF2 import PdfReader, PdfWriter
    import io

    pages = []

    with open(input_path, 'rb') as f:
        reader = PdfReader(f)

        for i, page in enumerate(reader.pages):
            writer = PdfWriter()
            writer.add_page(page)

            page_path = OUTPUT_DIR / f"page_{i + 1}.pdf"
            with open(page_path, 'wb') as pf:
                writer.write(pf)
            pages.append(page_path)

    return pages


async def _compress_pdf(input_path: Path, output_path: Path, level: str):
    """Compress PDF using PyPDF2 (basic)."""
    from PyPDF2 import PdfReader, PdfWriter

    reader = PdfReader(str(input_path))
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    # Simple compression by reducing quality would go here
    # For now, just copy
    with open(output_path, 'wb') as f:
        writer.write(f)


async def _convert_pdf_to_word(input_path: Path, output_path: Path, preserve_formatting: bool, include_images: bool):
    """Convert PDF to Word using python-docx (simplified)."""
    # In production, use pdf2docx or similar
    import docx
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = docx.Document()

    # Extract text from PDF (placeholder)
    # Real implementation would use pdfplumber to extract text + formatting
    doc.add_paragraph("Converted from PDF")
    doc.add_paragraph("This is a placeholder. Implement real conversion with pdfplumber + python-docx.")

    doc.save(output_path)


def _create_zip(file_paths: list[Path], zip_path: Path):
    """Create a ZIP archive from list of files."""
    import zipfile

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in file_paths:
            zipf.write(file_path, arcname=file_path.name)


@app.post("/api/tools/convert/word-to-pdf")
async def convert_word_to_pdf(
    file: UploadFile = File(...),
):
    """Convert Word document to PDF."""
    allowed_types = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File must be a Word document (DOC/DOCX)")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}.pdf"

    try:
        # Save uploaded file
        input_path = UPLOAD_DIR / f"{task_id}_input.docx"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # Placeholder conversion logic
        # In production, use docx2pdf or libreoffice
        await _convert_word_to_pdf(input_path, output_path)

        # Schedule auto-delete
        _schedule_auto_delete(output_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_path,
            filename="converted.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _convert_word_to_pdf(input_path: Path, output_path: Path):
    """Convert Word to PDF (placeholder)."""
    # In production, use docx2pdf or subprocess with libreoffice
    # For now, just create a dummy PDF
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(output_path), pagesize=letter)
    c.drawString(100, 750, "Converted from Word Document")
    c.drawString(100, 730, "This is a placeholder. Implement real conversion with docx2pdf or LibreOffice.")
    c.save()


@app.post("/api/tools/convert/pdf-to-txt")
async def convert_pdf_to_txt(
    file: UploadFile = File(...),
):
    """Convert PDF to plain text."""
    if file.content_type != 'application/pdf':
        raise HTTPException(status_code=400, detail="File must be a PDF")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}.txt"

    try:
        input_path = UPLOAD_DIR / f"{task_id}_input.pdf"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        await _convert_pdf_to_txt(input_path, output_path)

        _schedule_auto_delete(output_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_path,
            filename="extracted.txt",
            media_type="text/plain"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _convert_pdf_to_txt(input_path: Path, output_path: Path):
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    with pdfplumber.open(str(input_path)) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
        await f.write(text)


@app.post("/api/tools/convert/txt-to-pdf")
async def convert_txt_to_pdf(
    file: UploadFile = File(...),
):
    """Convert text file to PDF."""
    if file.content_type != 'text/plain':
        raise HTTPException(status_code=400, detail="File must be a TXT file")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}.pdf"

    try:
        input_path = UPLOAD_DIR / f"{task_id}_input.txt"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        await _convert_txt_to_pdf(input_path, output_path)

        _schedule_auto_delete(output_path)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_path,
            filename="converted.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _convert_txt_to_pdf(input_path: Path, output_path: Path):
    """Convert TXT to PDF (placeholder)."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(output_path), pagesize=letter)
    with open(input_path, 'r') as f:
        text = f.read()
    c.drawString(100, 750, "Converted from Text File")
    y = 730
    for line in text.split('\n')[:30]:
        c.drawString(100, y, line[:80])
        y -= 20
    c.save()


@app.post("/api/tools/convert/pdf-to-image")
async def convert_pdf_to_image(
    file: UploadFile = File(...),
    format: str = Form('png'),  # 'png' or 'jpeg'
    dpi: int = Form(200)
):
    """Convert PDF pages to images."""
    if file.content_type != 'application/pdf':
        raise HTTPException(status_code=400, detail="File must be a PDF")
    await validate_file_size(file, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_zip = OUTPUT_DIR / f"{task_id}_images.zip"

    try:
        input_path = UPLOAD_DIR / f"{task_id}_input.pdf"
        async with aiofiles.open(input_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        await _convert_pdf_to_image(input_path, output_zip, format, dpi)

        _schedule_auto_delete(output_zip)
        _schedule_auto_delete(input_path)

        return FileResponse(
            path=output_zip,
            filename="pdf_images.zip",
            media_type="application/zip"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _convert_pdf_to_image(input_path: Path, output_zip: Path, format: str, dpi: int):
    """Convert each PDF page to image and pack into ZIP."""
    from pdf2image import convert_from_path
    import zipfile
    import io

    # Convert PDF to list of images
    images = convert_from_path(
        str(input_path),
        dpi=dpi,
        fmt=format,
        thread_count=2
    )

    # Create ZIP in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for i, img in enumerate(images):
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format=format.upper())
            img_byte_arr.seek(0)
            zipf.writestr(f"page_{i+1}.{format}", img_byte_arr.read())

    zip_buffer.seek(0)
    async with aiofiles.open(output_zip, 'wb') as f:
        await f.write(zip_buffer.read())


@app.post("/api/tools/convert/image-to-pdf")
async def convert_image_to_pdf(
    files: list[UploadFile] = File(...),
):
    """Convert one or more images to PDF."""
    if len(files) == 0:
        raise HTTPException(status_code=400, detail="At least one image required")

    # Validate image types
    allowed = ['image/jpeg', 'image/png', 'image/webp']
    for f in files:
        if f.content_type not in allowed:
            raise HTTPException(status_code=400, detail=f"File {f.filename} is not a supported image (JPEG, PNG, WEBP)")
        await validate_file_size(f, MAX_FILE_SIZE)

    task_id = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{task_id}.pdf"

    try:
        image_paths = []
        for idx, file in enumerate(files):
            img_path = UPLOAD_DIR / f"{task_id}_img{idx}.{file.filename.split('.')[-1]}"
            async with aiofiles.open(img_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            image_paths.append(img_path)

        await _convert_image_to_pdf(image_paths, output_path)

        for p in image_paths:
            p.unlink(missing_ok=True)

        _schedule_auto_delete(output_path)

        return FileResponse(
            path=output_path,
            filename="converted.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _convert_image_to_pdf(image_paths: list[Path], output_path: Path):
    """Combine images into a single PDF using reportlab, each image on its own page."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from PIL import Image

    c = canvas.Canvas(str(output_path), pagesize=letter)
    page_width, page_height = letter

    for img_path in image_paths:
        # Open image to get dimensions
        img = Image.open(str(img_path))
        if img.mode in ('RGBA', 'LA'):
            # Convert to RGB (paste white background)
            rgb = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'RGBA':
                rgb.paste(img, mask=img.split()[-1])
            else:
                rgb.paste(img, mask=img)
            img = rgb

        # Compute scaled dimensions to fit within page margins
        max_width = page_width - 100  # leave 50pt margin each side
        max_height = page_height - 150  # leave margin top/bottom
        img_width, img_height = img.size

        # Convert pixel to points assuming 72 DPI (default for reportlab)
        # But we can just scale arbitrarily: we want to fit while preserving aspect ratio
        scale = min(max_width / img_width, max_height / img_height, 1.0)
        draw_width = img_width * scale
        draw_height = img_height * scale

        # Center position
        x = (page_width - draw_width) / 2
        y = (page_height - draw_height) / 2

        # Save image to a temporary file (PNG) to use as path
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img.save(tmp, format='PNG')
            tmp_path = tmp.name

        c.drawImage(tmp_path, x, y, width=draw_width, height=draw_height, preserveAspectRatio=True, mask='auto')
        # Clean up temp file
        import os
        os.unlink(tmp_path)

        c.showPage()

    c.save()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
